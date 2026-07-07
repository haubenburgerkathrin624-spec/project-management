from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pathlib import Path


OUT_DIR = Path("信息化项目预算与流程管理/改造方案输出")
DOCX_PATH = OUT_DIR / "信息化项目管理系统改造可行性与详细设计方案.docx"
MD_PATH = OUT_DIR / "信息化项目管理系统改造可行性与详细设计方案.md"


ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "EAF2F8"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def paragraph_text(paragraph):
    return "".join(run.text for run in paragraph.runs)


def add_para(doc, text="", style=None, bold=False, color=None, size=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        set_east_asia_font(run)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_east_asia_font(run)
        run.font.color.rgb = ACCENT if level < 3 else DARK
        run.font.bold = True
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_row_split(table.rows[0])
    set_table_width(table, [16.2])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_east_asia_font(r)
    r.bold = True
    r.font.color.rgb = DARK
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_east_asia_font(r2)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_row_split(table.rows[0])
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        set_cell_shading(hdr_cells[i], LIGHT_FILL)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_east_asia_font(r)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK
    for row in rows:
        tr = table.add_row()
        prevent_row_split(tr)
        cells = tr.cells
        for i, text in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            set_east_asia_font(r)
            r.font.size = Pt(9)
            r.font.color.rgb = INK
    set_table_width(table, widths_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("信息化项目管理系统改造方案")
    set_east_asia_font(r)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11.5, DARK),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)


def write_md():
    md = """# 信息化项目管理系统改造可行性与详细设计方案

## 1. 方案摘要

本方案基于《上海电机学院项目管理系统-用户手册》、两场需求沟通录音及已整理的《录音分析：旧系统痛点与改进建议》编制。旧系统已经具备项目管理、采购管理、实施、试运行、验收、归档、退金、结项、立项审核、验收审核、事务提醒、权限管理等基础能力，具备改造基础。后续改造不建议推倒重建，而应围绕高校信息化部门真实使用场景进行流程瘦身、预算联动、材料归集、查询看板和归档能力增强。

改造目标是把系统从“按阶段填表和审批的项目管理系统”，升级为“面向信息化项目预算与过程治理的轻量化工作台”。核心价值包括：让项目负责人少找模板、少重复填；让信息化部门少追材料、少人工汇总；让领导少审批、多看图；让归档审计少翻文件、多一键生成。

## 2. 旧系统现状

用户手册显示，旧系统已覆盖以下基础功能：

- 登录、系统管理、用户管理、部门管理、菜单权限控制。
- 项目管理：新建立项、团队角色配置、状态流转、条件搜索。
- 采购管理：新建采购信息、项目采购查询、采购信息编辑。
- 阶段管理：实施、试运行、验收、归档、退金、结项。
- 审批管理：立项审核、验收审核、审核通过/驳回。
- 事务提醒：待办和已处理事务筛选。

这些能力说明系统已经沉淀了项目主数据、角色权限、阶段状态、附件上传、审核意见和事务提醒等基础能力，技术上存在渐进式改造条件。

## 3. 用户痛点与设计回应

| 用户痛点 | 现象 | 改造回应 | 预期爽点 |
|---|---|---|---|
| 流程过重 | 采购、合同、审计等被设计得过细 | 复杂外部流程只做结果归集和材料留痕 | 不再被不必要审批卡住 |
| 查询不便 | 年度、预算、状态、执行情况难以快速汇总 | 增强组合筛选、项目台账、预算看板 | 一屏看到项目和预算全貌 |
| 归档费力 | 材料散落，后续还要人工追问 | 阶段材料清单、缺失提醒、一键归档清单 | 到归档时不用重新翻材料 |
| 领导负担大 | 立项、验收都让领导线上审批 | 领导侧改为看板和查询，必要节点线下签字留痕 | 领导看图不点流程 |
| 预算脱节 | 项目库和预算分离，立项决策不直观 | 项目库与年度预算同屏联动 | 像看购物车和银行卡一样做立项 |
| 模板缺失 | 申报、论证、试运行、验收材料格式不统一 | 建立信息化项目常用模板库 | 老师不再到处找模板 |
| 资料项不灵活 | 不同项目材料差异大，固定表单难适配 | 必填资料加自定义资料项 | 大小项目都能顺着走 |

## 4. 改造定位

系统定位为“高校信息化项目预算与流程管理系统”。不替代学校已有采购、合同、审计、档案系统，不追求大而全，而是服务信息化部门在项目全周期中的过程治理、预算决策、材料留痕和归档移交。

适用对象包括：信息化部门项目管理员、项目申报人/负责人、项目主管/归口负责人、部门领导、档案或综合管理人员、系统管理员。

## 5. 改造后主流程

推荐主流程为：项目申报 -> 项目库 -> 年度预算 -> 立项 -> 采购 -> 实施 -> 试运行 -> 验收 -> 归档/移交 -> 退金/结项。

设计原则：

- 申报全年开放，无批次限制。
- 项目库承载“想做的项目”，预算模块承载“能花的钱”，立项动作负责把二者匹配起来。
- 采购、实施等复杂环节不做完整审批替代，只做过程资料、结果文件、合同等材料归集。
- 领导不作为普通流程审批节点，主要查看报表、可视化图和项目详情。
- 归档以材料完整性为核心，系统自动生成归档清单和缺失提醒。

## 6. 功能模块详细设计

### 6.1 首页工作台

首页应按角色显示不同内容：

- 项目负责人：我的项目、待补材料、被退回事项、即将到期节点。
- 信息化部门管理员：待审核/待入库/待立项项目、预算余额、材料缺失项目、验收和归档预警。
- 领导：年度项目概览、预算执行、项目阶段分布、重点项目进展、异常提醒。

爽点：登录后不用先找菜单，系统直接告诉用户“今天该处理什么、哪些项目卡住了、钱还剩多少”。

### 6.2 项目申报

保留旧系统新建项目能力，调整为全年开放申报。支持草稿保存、附件上传、退回修改、版本留痕和模板下载。

建议字段包括：项目名称、申报部门、项目负责人、建设背景、建设目标、需求说明、预估金额、预期成果、紧急程度、建议年度、附件材料。

### 6.3 项目库

项目库用于沉淀已通过初步审核或评审的储备项目。支持项目状态、优先级、建议年度、预算估算、评审结果、入库时间、是否已立项等信息。

关键设计：项目库不等于已立项，项目入库后可以等待预算安排；立项时再从项目库中选择项目并绑定预算。

### 6.4 年度预算

新增年度预算模块。预算信息应至少包括预算年度、预算名称、预算编号、预算金额、预算性质、预算来源、已占用金额、已执行金额、余额、关联项目。

预算模块应支持导入、调整、冻结、释放、占用、执行登记和预算台账导出。立项时自动进行预算占用，项目取消或调整时可释放预算。

### 6.5 立项管理

立项不再从零开始填写全部内容，而是从项目库选择候选项目，结合年度预算做决策。页面建议采用“左侧项目库，右侧预算可用情况”的同屏设计。

立项通过后生成正式项目编号、项目状态、预算占用记录和后续阶段任务。

### 6.6 采购管理

采购环节保持轻量化。系统不替代学校采购系统，只记录采购过程文件、结果文件和合同信息。

建议资料项包括：采购方式、采购金额、论证材料、采购过程文件、采购结果文件、中标/成交信息、合同文件、合同金额、供应商信息。

不同金额和采购方式只作为材料提示和统计字段，不展开成复杂流程。

### 6.7 实施管理

实施阶段重点收集过程资料。系统预置若干常见资料项，例如启动会材料、实施计划、会议纪要、阶段成果、培训材料、测评材料等，同时允许项目负责人自定义追加资料。

爽点：系统给出“建议上传清单”，但不把项目卡死在固定模板里。

### 6.8 试运行管理

试运行阶段应支持试运行计划、培训记录、问题整改、试运行报告、验收准备情况等信息。系统可提供试运行报告模板，并在试运行报告上传后提示是否具备验收申请条件。

### 6.9 验收管理

保留旧系统验收申请和验收审核能力，但简化领导审批。验收重点记录验收申请、验收报告、验收意见、整改记录、验收结论和必要签字材料。

线上审核建议由信息化部门管理员或项目主管处理，领导签字材料可以作为附件留痕。

### 6.10 归档/移交

归档模块是本次改造的关键。系统应按项目阶段自动汇总已上传材料，生成归档清单，标记缺失材料，并支持一键打包下载。

建议支持：

- 按阶段生成材料目录。
- 自动识别必填材料是否缺失。
- 手动补充线下签字件。
- 生成归档清单 Word/PDF。
- 批量移交档案馆或综合管理部门。
- 归档状态跟踪。

### 6.11 模板中心

新增“文档模板库”。模板按流程节点分类，包括项目申报书、采购论证材料、预算说明、实施过程资料、试运行报告、验收申请、验收报告、归档清单等。

模板应支持版本管理、启停、按项目类型推荐和下载记录。

### 6.12 领导看板

领导侧以只读为主。看板包括年度项目数量、预算总额、已占用/已执行/剩余额度、项目阶段分布、超期项目、待验收项目、已归档项目、重点项目详情。

爽点：领导不用进每个审批页面，只要看图、查项目、看异常。

## 7. 权限与角色设计

| 角色 | 主要职责 | 系统权限 |
|---|---|---|
| 项目申报人/负责人 | 申报项目、补充资料、提交验收 | 新建申报、编辑本人项目、上传资料、查看本人项目进度 |
| 信息化部门管理员 | 审核、入库、立项、预算匹配、推进归档 | 全量项目管理、预算管理、材料清单、归档确认、统计导出 |
| 项目主管/归口负责人 | 专业确认和必要节点审核 | 查看负责项目、审核关键节点、填写意见 |
| 领导 | 掌握全局和重点项目 | 查看看板、查询项目、导出报表，不参与普通细节审批 |
| 档案/综合管理员 | 接收归档材料 | 查看归档包、确认移交、下载清单 |
| 系统管理员 | 账号、角色、菜单、参数配置 | 用户、部门、菜单、流程、模板、字典配置 |

## 8. 数据设计

核心数据对象包括：项目主表、项目库表、年度预算表、预算占用表、采购信息表、阶段材料表、模板表、审核记录表、待办任务表、归档包表、操作日志表。

数据改造重点是新增预算和归档两条主线：预算要能与项目立项绑定，归档要能与各阶段附件绑定。

## 9. 技术可行性

### 9.1 可复用能力

旧系统已有登录、权限、项目管理、采购管理、阶段状态、审核、附件上传、事务提醒等能力，可作为改造基础。

### 9.2 新增能力

需要新增年度预算、项目库、模板中心、归档清单、领导看板、材料缺失提醒、大附件支持、预算占用释放等能力。

### 9.3 集成边界

近期不强制对接采购、合同、审计、档案系统；先通过结果文件上传和归档清单满足管理需要。后续可通过接口逐步对接统一身份认证、财务预算、采购平台、档案系统。

### 9.4 非功能要求

- 附件：建议支持单文件 200MB，提供上传进度和失败重传。
- 安全：按角色、部门、项目团队控制数据可见范围。
- 审计：保留关键操作日志、审核意见、附件版本。
- 性能：常用查询 3 秒内返回，统计看板支持年度数据秒级展示。
- 备份：项目资料和附件需纳入定期备份。

## 10. 实施路径

| 阶段 | 时间建议 | 主要工作 | 交付物 |
|---|---|---|---|
| 0. 需求确认 | 1-2 周 | 确认流程边界、角色、字段、材料清单、预算口径 | 需求确认稿、原型范围 |
| 1. 基础改造 | 3-4 周 | 调整主流程、项目库、角色权限、查询列表 | 可演示版本 1 |
| 2. 预算联动 | 4-5 周 | 年度预算、预算占用、立项同屏决策 | 预算立项版本 |
| 3. 材料与归档 | 4-5 周 | 模板中心、材料清单、缺失提醒、一键归档 | 归档闭环版本 |
| 4. 看板与试点 | 3-4 周 | 领导看板、统计报表、试点数据迁移、问题优化 | 试运行版本 |
| 5. 上线推广 | 2 周 | 培训、初始化数据、上线保障 | 正式上线版本 |

## 11. 风险与对策

| 风险 | 影响 | 对策 |
|---|---|---|
| 需求继续扩张 | 系统再次变复杂 | 坚持“过程留痕和结果归集”，复杂外部流程只做附件化 |
| 历史数据不规范 | 迁移和统计困难 | 先迁移核心字段，历史附件按项目归档，不强求全量结构化 |
| 预算口径不统一 | 预算统计不可信 | 明确预算编号、性质、金额、占用、执行、释放规则 |
| 用户不愿填材料 | 归档仍然缺失 | 模板、缺失提醒、阶段必填项和待办推送结合 |
| 领导审批习惯差异 | 流程难统一 | 领导默认只读，确需审批的节点做可配置 |

## 12. 验收指标

- 项目申报、项目库、预算、立项、采购、实施、试运行、验收、归档主流程可闭环。
- 年度预算可查看总额、已占用、已执行和余额。
- 立项时可从项目库选择项目并绑定预算。
- 每个阶段可上传材料，并能区分必填、选填和自定义材料。
- 可自动生成归档清单和项目资料包。
- 领导可查看年度项目和预算看板，无需参与普通审批。
- 常用查询支持年度、状态、负责人、部门、金额、归档状态等条件。
- 模板中心可按阶段维护并下载模板。

## 13. 结论

本次改造具备业务可行性、技术可行性和实施可控性。建议采用“保留旧系统基础能力 + 新增预算与归档核心能力 + 简化复杂审批流程”的路线。这样既避免推倒重来，也能直接回应老师提出的真实痛点，并形成明显的使用爽点：少填、少审、少找、好查、好归档、领导少打扰、部门好管理。
"""
    MD_PATH.write_text(md, encoding="utf-8")
    return md


def build_docx():
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("可行性与详细设计方案")
    set_east_asia_font(r)
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    r.bold = True

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.05
    r = title.add_run("信息化项目管理系统改造")
    set_east_asia_font(r)
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    title2 = doc.add_paragraph()
    title2.paragraph_format.space_after = Pt(14)
    r2 = title2.add_run("面向高校信息化部门的预算、流程与归档一体化升级")
    set_east_asia_font(r2)
    r2.font.size = Pt(13)
    r2.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("编制依据", "用户手册、两场需求沟通录音、录音痛点分析"),
        ("适用对象", "信息化部门、项目负责人、项目主管、领导、档案/综合管理人员"),
        ("方案定位", "在旧系统基础上渐进式改造，不推倒重建"),
        ("核心目标", "少填、少审、少找、好查、好归档"),
    ]
    for i, (k, v) in enumerate(rows):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        set_cell_shading(meta.cell(i, 0), LIGHT_FILL)
        for c in meta.row_cells(i):
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for rr in p.runs:
                    set_east_asia_font(rr)
                    rr.font.size = Pt(9.5)
        meta.cell(i, 0).paragraphs[0].runs[0].bold = True
    set_table_width(meta, [3.0, 13.2])
    add_callout(
        doc,
        "总体结论",
        "建议采用“保留旧系统基础能力 + 新增预算与归档核心能力 + 简化复杂审批流程”的改造路线。该路线能直接回应老师提出的真实痛点，并把系统从按阶段填表的项目管理工具升级为信息化部门的项目治理工作台。",
    )

    add_heading(doc, "1. 方案摘要", 1)
    add_para(
        doc,
        "本方案基于《上海电机学院项目管理系统-用户手册》、两场需求沟通录音及已整理的痛点分析编制。旧系统已经具备项目管理、采购管理、实施、试运行、验收、归档、退金、结项、立项审核、验收审核、事务提醒和权限管理等基础能力，具备渐进式改造条件。",
    )
    add_para(
        doc,
        "改造目标是把系统从“按阶段填表和审批的项目管理系统”，升级为“面向信息化项目预算与过程治理的轻量化工作台”。核心价值包括：项目负责人少找模板、少重复填；信息化部门少追材料、少人工汇总；领导少审批、多看图；归档审计少翻文件、多一键生成。",
    )

    add_heading(doc, "2. 旧系统现状与可复用能力", 1)
    add_para(doc, "用户手册显示，旧系统已覆盖以下基础功能：")
    for item in [
        "登录、系统管理、用户管理、部门管理、菜单权限控制。",
        "项目管理：新建立项、团队角色配置、状态流转、条件搜索。",
        "采购管理：新建采购信息、项目采购查询、采购信息编辑。",
        "阶段管理：实施、试运行、验收、归档、退金、结项。",
        "审批管理：立项审核、验收审核、审核通过/驳回。",
        "事务提醒：待办和已处理事务筛选。",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "这些能力说明系统已经沉淀了项目主数据、角色权限、阶段状态、附件上传、审核意见和事务提醒等基础能力。后续改造应尽量复用这些基础设施，把投入集中在预算联动、项目库、模板中心、归档清单和领导看板上。",
    )

    add_heading(doc, "3. 用户痛点与设计回应", 1)
    add_table(
        doc,
        ["用户痛点", "现象", "改造回应", "预期爽点"],
        [
            ["流程过重", "采购、合同、审计等被设计得过细", "复杂外部流程只做结果归集和材料留痕", "不再被不必要审批卡住"],
            ["查询不便", "年度、预算、状态、执行情况难以快速汇总", "增强组合筛选、项目台账、预算看板", "一屏看到项目和预算全貌"],
            ["归档费力", "材料散落，后续还要人工追问", "阶段材料清单、缺失提醒、一键归档清单", "到归档时不用重新翻材料"],
            ["领导负担大", "立项、验收都让领导线上审批", "领导侧改为看板和查询，必要节点线下签字留痕", "领导看图不点流程"],
            ["预算脱节", "项目库和预算分离，立项决策不直观", "项目库与年度预算同屏联动", "像看购物车和银行卡一样做立项"],
            ["模板缺失", "申报、论证、试运行、验收材料格式不统一", "建立信息化项目常用模板库", "老师不再到处找模板"],
            ["资料项不灵活", "不同项目材料差异大，固定表单难适配", "必填资料加自定义资料项", "大小项目都能顺着走"],
        ],
        [3.0, 4.2, 5.2, 3.8],
    )

    add_heading(doc, "4. 改造定位与边界", 1)
    add_para(
        doc,
        "系统定位为“高校信息化项目预算与流程管理系统”。它不替代学校已有采购、合同、审计、档案系统，也不追求大而全，而是服务信息化部门在项目全周期中的过程治理、预算决策、材料留痕和归档移交。",
    )
    add_callout(
        doc,
        "边界原则",
        "外部已有成熟系统承载的采购、合同、审计等复杂流程，本系统只保留必要字段、过程文件和结果文件。这样既能合规留痕，又不会把系统做成另一个难维护的大平台。",
    )

    add_heading(doc, "5. 改造后主流程", 1)
    add_para(doc, "推荐主流程为：")
    for item in [
        "项目申报：全年开放，无批次限制。",
        "项目库：承载想做、待安排、待预算匹配的储备项目。",
        "年度预算：管理预算编号、金额、性质、占用、执行和余额。",
        "立项：从项目库选择项目，并与预算同屏匹配。",
        "采购：记录采购过程文件、结果文件、合同及供应商信息。",
        "实施：收集过程资料，支持自定义资料项。",
        "试运行：沉淀培训、问题整改、试运行报告和验收准备。",
        "验收：提交验收申请、验收报告、结论和签字材料。",
        "归档/移交：自动生成材料清单、归档包和移交状态。",
        "退金/结项：保留旧系统已有闭环，可按学校流程配置。",
    ]:
        add_number(doc, item)

    add_heading(doc, "6. 功能模块详细设计", 1)
    modules = [
        ("6.1 首页工作台", [
            "项目负责人看到我的项目、待补材料、被退回事项和即将到期节点。",
            "信息化部门管理员看到待审核、待入库、待立项项目，预算余额，材料缺失项目，验收和归档预警。",
            "领导看到年度项目概览、预算执行、项目阶段分布、重点项目进展和异常提醒。",
            "爽点：登录后不用找菜单，系统直接告诉用户今天该处理什么、哪些项目卡住了、钱还剩多少。",
        ]),
        ("6.2 项目申报", [
            "保留旧系统新建项目能力，调整为全年开放申报。",
            "支持草稿保存、附件上传、退回修改、版本留痕和模板下载。",
            "建议字段包括项目名称、申报部门、项目负责人、建设背景、建设目标、需求说明、预估金额、预期成果、紧急程度、建议年度、附件材料。",
        ]),
        ("6.3 项目库", [
            "项目库用于沉淀已通过初步审核或评审的储备项目。",
            "支持项目状态、优先级、建议年度、预算估算、评审结果、入库时间、是否已立项等信息。",
            "项目库不等于已立项，项目入库后可以等待预算安排；立项时再从项目库中选择项目并绑定预算。",
        ]),
        ("6.4 年度预算", [
            "新增年度预算模块，管理预算年度、预算名称、预算编号、预算金额、预算性质、预算来源、已占用金额、已执行金额、余额和关联项目。",
            "支持导入、调整、冻结、释放、占用、执行登记和预算台账导出。",
            "立项时自动进行预算占用，项目取消或调整时可释放预算。",
        ]),
        ("6.5 立项管理", [
            "立项不再从零开始填写全部内容，而是从项目库选择候选项目，结合年度预算做决策。",
            "页面建议采用“左侧项目库，右侧预算可用情况”的同屏设计。",
            "立项通过后生成正式项目编号、项目状态、预算占用记录和后续阶段任务。",
        ]),
        ("6.6 采购管理", [
            "采购环节保持轻量化，不替代学校采购系统。",
            "建议资料项包括采购方式、采购金额、论证材料、采购过程文件、采购结果文件、中标/成交信息、合同文件、合同金额、供应商信息。",
            "不同金额和采购方式只作为材料提示和统计字段，不展开成复杂流程。",
        ]),
        ("6.7 实施管理", [
            "实施阶段重点收集过程资料。",
            "系统预置启动会材料、实施计划、会议纪要、阶段成果、培训材料、测评材料等常见资料项。",
            "允许项目负责人自定义追加资料，兼容大小项目差异。",
        ]),
        ("6.8 试运行管理", [
            "支持试运行计划、培训记录、问题整改、试运行报告、验收准备情况等信息。",
            "提供试运行报告模板，并在试运行报告上传后提示是否具备验收申请条件。",
        ]),
        ("6.9 验收管理", [
            "保留旧系统验收申请和验收审核能力，但简化领导审批。",
            "重点记录验收申请、验收报告、验收意见、整改记录、验收结论和必要签字材料。",
            "线上审核建议由信息化部门管理员或项目主管处理，领导签字材料作为附件留痕。",
        ]),
        ("6.10 归档/移交", [
            "按项目阶段自动汇总已上传材料，生成归档清单，标记缺失材料，并支持一键打包下载。",
            "支持手动补充线下签字件、批量移交档案馆或综合管理部门、归档状态跟踪。",
            "爽点：到归档时不再重新翻文件，系统直接告诉你缺什么、能不能移交。",
        ]),
        ("6.11 模板中心", [
            "新增文档模板库，按流程节点分类。",
            "模板包括项目申报书、采购论证材料、预算说明、实施过程资料、试运行报告、验收申请、验收报告、归档清单等。",
            "支持版本管理、启停、按项目类型推荐和下载记录。",
        ]),
        ("6.12 领导看板", [
            "领导侧以只读为主。",
            "看板包括年度项目数量、预算总额、已占用/已执行/剩余额度、项目阶段分布、超期项目、待验收项目、已归档项目、重点项目详情。",
            "爽点：领导不用进每个审批页面，只要看图、查项目、看异常。",
        ]),
    ]
    for heading, bullets in modules:
        add_heading(doc, heading, 2)
        for bullet in bullets:
            add_bullet(doc, bullet)

    add_heading(doc, "7. 权限与角色设计", 1)
    add_table(
        doc,
        ["角色", "主要职责", "系统权限"],
        [
            ["项目申报人/负责人", "申报项目、补充资料、提交验收", "新建申报、编辑本人项目、上传资料、查看本人项目进度"],
            ["信息化部门管理员", "审核、入库、立项、预算匹配、推进归档", "全量项目管理、预算管理、材料清单、归档确认、统计导出"],
            ["项目主管/归口负责人", "专业确认和必要节点审核", "查看负责项目、审核关键节点、填写意见"],
            ["领导", "掌握全局和重点项目", "查看看板、查询项目、导出报表，不参与普通细节审批"],
            ["档案/综合管理员", "接收归档材料", "查看归档包、确认移交、下载清单"],
            ["系统管理员", "账号、角色、菜单、参数配置", "用户、部门、菜单、流程、模板、字典配置"],
        ],
        [3.2, 5.2, 7.8],
    )

    add_heading(doc, "8. 数据设计", 1)
    add_para(doc, "核心数据对象建议如下：")
    add_table(
        doc,
        ["数据对象", "关键字段", "说明"],
        [
            ["项目主表", "项目编号、名称、部门、负责人、状态、年度、类型", "承载正式项目信息和状态流转"],
            ["项目库表", "入库编号、评审结果、优先级、建议年度、是否立项", "承载储备项目"],
            ["年度预算表", "预算编号、名称、金额、性质、来源、年度、余额", "承载年度预算池"],
            ["预算占用表", "项目、预算、占用金额、释放金额、执行金额", "连接项目与预算"],
            ["阶段材料表", "项目、阶段、材料名称、必填标识、附件、版本", "支撑过程资料和归档"],
            ["模板表", "模板名称、阶段、版本、启停状态、适用类型", "支撑模板中心"],
            ["归档包表", "项目、清单、缺失项、打包文件、移交状态", "支撑归档和移交"],
            ["操作日志表", "用户、动作、时间、对象、前后值", "支撑审计追溯"],
        ],
        [3.0, 6.2, 7.0],
    )

    add_heading(doc, "9. 技术可行性", 1)
    add_heading(doc, "9.1 可复用能力", 2)
    add_para(doc, "旧系统已有登录、权限、项目管理、采购管理、阶段状态、审核、附件上传、事务提醒等能力，可作为改造基础。")
    add_heading(doc, "9.2 新增能力", 2)
    for item in [
        "年度预算、预算占用、预算释放与执行登记。",
        "项目库和立项同屏决策。",
        "模板中心、阶段材料清单、缺失材料提醒。",
        "归档清单生成、项目资料包打包下载。",
        "领导看板和统计报表。",
        "大附件上传能力，建议单文件支持到 200MB。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "9.3 集成边界", 2)
    add_para(doc, "近期不强制对接采购、合同、审计、档案系统；先通过结果文件上传和归档清单满足管理需要。后续可通过接口逐步对接统一身份认证、财务预算、采购平台和档案系统。")
    add_heading(doc, "9.4 非功能要求", 2)
    for item in [
        "附件：支持上传进度、失败重传、文件类型校验和容量控制。",
        "安全：按角色、部门、项目团队控制数据可见范围。",
        "审计：保留关键操作日志、审核意见、附件版本。",
        "性能：常用查询 3 秒内返回，统计看板支持年度数据快速展示。",
        "备份：项目资料和附件纳入定期备份。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. 实施路径", 1)
    add_table(
        doc,
        ["阶段", "时间建议", "主要工作", "交付物"],
        [
            ["0. 需求确认", "1-2 周", "确认流程边界、角色、字段、材料清单、预算口径", "需求确认稿、原型范围"],
            ["1. 基础改造", "3-4 周", "调整主流程、项目库、角色权限、查询列表", "可演示版本 1"],
            ["2. 预算联动", "4-5 周", "年度预算、预算占用、立项同屏决策", "预算立项版本"],
            ["3. 材料与归档", "4-5 周", "模板中心、材料清单、缺失提醒、一键归档", "归档闭环版本"],
            ["4. 看板与试点", "3-4 周", "领导看板、统计报表、试点数据迁移、问题优化", "试运行版本"],
            ["5. 上线推广", "2 周", "培训、初始化数据、上线保障", "正式上线版本"],
        ],
        [2.6, 2.6, 7.0, 4.0],
    )

    add_heading(doc, "11. 风险与对策", 1)
    add_table(
        doc,
        ["风险", "影响", "对策"],
        [
            ["需求继续扩张", "系统再次变复杂", "坚持过程留痕和结果归集，复杂外部流程只做附件化"],
            ["历史数据不规范", "迁移和统计困难", "先迁移核心字段，历史附件按项目归档，不强求全量结构化"],
            ["预算口径不统一", "预算统计不可信", "明确预算编号、性质、金额、占用、执行、释放规则"],
            ["用户不愿填材料", "归档仍然缺失", "模板、缺失提醒、阶段必填项和待办推送结合"],
            ["领导审批习惯差异", "流程难统一", "领导默认只读，确需审批的节点做可配置"],
        ],
        [4.0, 4.0, 8.2],
    )

    add_heading(doc, "12. 验收指标", 1)
    for item in [
        "项目申报、项目库、预算、立项、采购、实施、试运行、验收、归档主流程可闭环。",
        "年度预算可查看总额、已占用、已执行和余额。",
        "立项时可从项目库选择项目并绑定预算。",
        "每个阶段可上传材料，并能区分必填、选填和自定义材料。",
        "可自动生成归档清单和项目资料包。",
        "领导可查看年度项目和预算看板，无需参与普通审批。",
        "常用查询支持年度、状态、负责人、部门、金额、归档状态等条件。",
        "模板中心可按阶段维护并下载模板。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "13. 结论", 1)
    add_para(
        doc,
        "本次改造具备业务可行性、技术可行性和实施可控性。建议采用“保留旧系统基础能力 + 新增预算与归档核心能力 + 简化复杂审批流程”的路线。这样既避免推倒重来，也能直接回应老师提出的真实痛点，并形成明显的使用爽点：少填、少审、少找、好查、好归档、领导少打扰、部门好管理。",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    write_md()
    build_docx()
    print(DOCX_PATH)
    print(MD_PATH)
